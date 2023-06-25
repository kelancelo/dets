from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from .models import *
from django.core.exceptions import ObjectDoesNotExist
from django.core import serializers
from django.http import HttpResponseNotFound, JsonResponse, HttpResponse
from django.contrib.auth import logout
from django.contrib.auth.models import Group
from django.shortcuts import get_object_or_404
from docx import Document
import os


@login_required
def index(request):
    context = {}
    # Check if there's an active event
    event = None
    try:
        event = Event.objects.get(status="active")
    except ObjectDoesNotExist:
        return render(request, "main/index.html")
    
    context["event"] = event
    # If there's an active event, check if the logged in judge is a judge of that event
    is_judge_of_event = event.judges.filter(id=request.user.id).exists()
    context["is_judge_of_event"] = is_judge_of_event
    # Get the rounds of the event
    rounds = event.round_set.all()
    context["rounds"] = rounds
    
    return render(request, "main/index.html", context)


# A function for checking if a judge has already given a participant the scores
# for the criteria of the active round
@login_required
def check_if_participant_has_been_scored(request, participant_id, round_id):
    if Score.objects.filter(
        judge=request.user,
        participant_id=participant_id,
        criterion__round__id=round_id
    ).exists():
        scores = Score.objects.filter(
            judge=request.user,
            participant_id=participant_id,
            criterion__round__id=round_id
        )
        return JsonResponse(serializers.serialize("json", scores), safe=False)
    else:
        return HttpResponseNotFound()


# Save participant's scores in the database
@login_required 
def save_scores(request):
    participant_id = request.POST["participant_id"]
    for key, value in request.POST.items():
        if "criterion" in key:
            criterion_id = key.split("-")[1]
            Score.objects.create(
                criterion_id=criterion_id, 
                judge=request.user, 
                participant_id=participant_id,
                value=value
            )
    return redirect("main:index")
    

# Returns a list of the participants of the active round
@login_required
def get_participants(request, round_id):
    participants = Round.objects.get(id=round_id).participants.all()
    return JsonResponse(serializers.serialize("json", participants), safe=False)


@login_required
def get_extra_participant_data(request, round_id, participant_id):
    event_participant = Event_Participant.objects.get(
        event=Round.objects.get(id=round_id).event,
        participant_id=participant_id
    )
    return JsonResponse({
        "candidate_number": event_participant.candidate_number,
        "represented_entity": event_participant.represented_entity.name
    }, safe=False)


@login_required
def get_criteria(request, round_id):
    criteria = Criterion.objects.filter(round_id=round_id)
    return JsonResponse(serializers.serialize("json", criteria), safe=False)


@login_required
def custom_logout(request):
    logout(request)
    return redirect("main:login")


def get_last_event_info(request):
    event = Event.objects.filter(status="ended").order_by("-end_date").first()
    final_round = get_object_or_404(Round, event=event, num_winners=1)
    total_scores = Total_Score.objects.filter(round=final_round).order_by("-value")[:3]
    winners = []
    for ts in total_scores:
        participant = ts.participant
        winners.append({
            "candidate_number": Event_Participant.objects.get(event=event, participant=participant).candidate_number, 
            "name": f"{participant.first_name} {participant.middle_name if participant.middle_name else ''} {participant.last_name}", 
            "total_score": ts.value
        })
    return JsonResponse({
        "event_name": event.name,
        "winners": winners
    }, safe=False)


def int_to_ordinal(num):
    suffixes = {1: 'st', 2: 'nd', 3: 'rd'}
    if 10 < num < 20:
        suffix = 'th'
    else:
        suffix = suffixes.get(num % 10, 'th')
    return str(num) + suffix


def download_last_event_info(request):
    event = Event.objects.filter(status="ended").order_by("-end_date").first()
    final_round = get_object_or_404(Round, event=event, num_winners=1)
    total_scores = Total_Score.objects.filter(round=final_round).order_by("-value")[:3]
    filename = event.name + ".docx"
    document = Document()
    document.add_paragraph().add_run()
    document.add_heading(event.name, level=1)
    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[1].text = "Candidate number"
    hdr_cells[2].text = "Name"
    hdr_cells[3].text = "Total score"
    for i, ts in enumerate(total_scores):
        participant = ts.participant
        row_cells = table.add_row().cells
        row_cells[0].text = "Winner" if i == 0 else int_to_ordinal(i) + " runner up"
        row_cells[1].text = str(Event_Participant.objects.get(event=event, participant=participant).candidate_number)
        row_cells[2].text = f"{participant.first_name} {participant.middle_name if participant.middle_name else ''} {participant.last_name}"
        row_cells[3].text = str(ts.value) 
    document.add_page_break()
    document.save(filename)
    file = open(filename, "rb")
    response = HttpResponse(file.read(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response["Content-Disposition"] = f'attachment; filename={filename}'
    file.close()
    os.remove(filename)
    return response


def get_stats(request):
    users = list(User.objects.all().values("groups__name"))
    groups = list(Group.objects.all().values())
    events = list(Event.objects.all().values())
    participants = list(Participant.objects.all().values())
 
    return JsonResponse({
        "users": users,
        "groups": groups,
        "events": events,
        "participants": participants
    }, safe=False)


